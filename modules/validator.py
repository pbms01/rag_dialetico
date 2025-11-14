"""
═══════════════════════════════════════════════════════════════════════════
VALIDADOR E FORMATADOR DE CONTESTAÇÕES
═══════════════════════════════════════════════════════════════════════════
Valida qualidade e formata para DOCX
"""

import re
from typing import Dict, List, Tuple
from pathlib import Path
from datetime import datetime
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import Config

class ValidadorContestacao:
    """Valida qualidade da contestação gerada"""
    
    def __init__(self):
        pass
    
    def validar(self, contestacao: str) -> Dict:
        """
        Valida contestação e retorna métricas de qualidade
        
        Args:
            contestacao: Texto da contestação
            
        Returns:
            Dict com métricas e alertas
        """
        metricas = {}
        alertas = []
        
        # 1. Validar tamanho
        tamanho = len(contestacao)
        metricas['tamanho_caracteres'] = tamanho
        
        if tamanho < Config.MIN_CONTESTACAO_LENGTH:
            alertas.append(f"⚠️  Contestação muito curta ({tamanho} caracteres)")
        elif tamanho > Config.MAX_CONTESTACAO_LENGTH:
            alertas.append(f"⚠️  Contestação muito longa ({tamanho} caracteres)")
        
        # 2. Verificar presença de seções obrigatórias
        secoes_presentes = []
        secoes_faltantes = []
        
        for secao in Config.SECOES_OBRIGATORIAS:
            # Buscar seção (case insensitive, com variações)
            padrao = secao.replace(' ', r'\s+')
            if re.search(padrao, contestacao, re.IGNORECASE):
                secoes_presentes.append(secao)
            else:
                secoes_faltantes.append(secao)
        
        metricas['completude_estrutural'] = len(secoes_presentes) / len(Config.SECOES_OBRIGATORIAS)
        
        if secoes_faltantes:
            alertas.append(f"⚠️  Seções faltantes: {', '.join(secoes_faltantes)}")
        
        # 3. Detectar citações legais
        citacoes_lei = re.findall(
            r'(?:art\.|artigo)\s*\d+[º°]?(?:[-,]\s*§\s*\d+[º°]?)?.*?(?:Lei|CF|CDC)',
            contestacao,
            re.IGNORECASE
        )
        metricas['citacoes_legais'] = len(citacoes_lei)
        
        if len(citacoes_lei) < 3:
            alertas.append("⚠️  Poucas citações legais detectadas")
        
        # 4. Detectar precedentes/jurisprudência
        precedentes = re.findall(
            r'(?:jurisprudência|precedente|acórdão|súmula|STJ|STF|TJRJ)',
            contestacao,
            re.IGNORECASE
        )
        metricas['mencoes_jurisprudencia'] = len(precedentes)
        
        # 5. Densidade de fundamentação
        tokens = contestacao.split()
        metricas['densidade_fundamentacao'] = len(citacoes_lei) / len(tokens) if tokens else 0
        
        # 6. Estrutura de argumentação (presença de conectivos lógicos)
        conectivos = [
            'portanto', 'assim', 'dessa forma', 'consequentemente',
            'ademais', 'além disso', 'outrossim', 'por outro lado',
            'em que pese', 'não obstante', 'contudo', 'todavia'
        ]
        
        count_conectivos = sum(1 for c in conectivos if c in contestacao.lower())
        metricas['conectivos_argumentativos'] = count_conectivos
        
        # 7. Score geral de qualidade (0-100)
        score = 0
        score += min(metricas['completude_estrutural'] * 40, 40)  # 40% - Estrutura
        score += min((metricas['citacoes_legais'] / 5) * 30, 30)  # 30% - Citações
        score += min((metricas['mencoes_jurisprudencia'] / 3) * 15, 15)  # 15% - Jurisprudência
        score += min((count_conectivos / 5) * 15, 15)  # 15% - Argumentação
        
        metricas['score_qualidade'] = round(score, 1)
        
        # Classificação
        if score >= 80:
            classificacao = "Excelente"
        elif score >= 60:
            classificacao = "Boa"
        elif score >= 40:
            classificacao = "Regular"
        else:
            classificacao = "Necessita Revisão"
        
        metricas['classificacao'] = classificacao
        
        return {
            'metricas': metricas,
            'alertas': alertas,
            'valido': len(alertas) == 0 or metricas['score_qualidade'] >= 50
        }


class FormatadorDOCX:
    """Formata contestação em documento DOCX profissional"""
    
    def __init__(self):
        pass
    
    def criar_docx(
        self,
        contestacao: str,
        metadados: Dict,
        output_path: Path
    ) -> Path:
        """
        Cria documento DOCX formatado
        
        Args:
            contestacao: Texto da contestação
            metadados: Metadados da geração
            output_path: Caminho de saída
            
        Returns:
            Path do arquivo criado
        """
        doc = docx.Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)
        
        # Cabeçalho
        self._adicionar_cabecalho(doc, metadados)
        
        # Conteúdo da contestação
        self._adicionar_conteudo(doc, contestacao)
        
        # Rodapé
        self._adicionar_rodape(doc, metadados)
        
        # Salvar
        doc.save(output_path)
        print(f"✅ Documento DOCX salvo: {output_path}")
        
        return output_path
    
    def _adicionar_cabecalho(self, doc, metadados):
        """Adiciona cabeçalho ao documento"""
        # Título centralizado
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("CONTESTAÇÃO")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph()  # Espaço
        
        # Identificação
        p = doc.add_paragraph()
        p.add_run("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA ___ª VARA CÍVEL DA COMARCA DO RIO DE JANEIRO").bold = True
        
        doc.add_paragraph()  # Espaço
        
        # Informações sobre geração (comentário discreto)
        info = doc.add_paragraph()
        run = info.add_run(f"[Gerado automaticamente em {datetime.now().strftime('%d/%m/%Y %H:%M')}]")
        run.font.size = Pt(8)
        run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Espaço
    
    def _adicionar_conteudo(self, doc, contestacao):
        """Adiciona conteúdo formatado"""
        # Dividir por seções principais
        secoes = re.split(r'\n(?=#{1,3}\s|\d+\.\s[A-Z])', contestacao)
        
        for secao in secoes:
            if not secao.strip():
                continue
            
            # Detectar se é título principal
            if secao.startswith('#'):
                nivel = secao.count('#')
                texto = secao.lstrip('#').strip()
                
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.bold = True
                
                if nivel == 1:
                    run.font.size = Pt(13)
                elif nivel == 2:
                    run.font.size = Pt(12)
                else:
                    run.font.size = Pt(11)
                
                doc.add_paragraph()  # Espaço após título
            
            # Detectar numeração (ex: 1., 1.1., a))
            elif re.match(r'^\d+\.', secao.strip()):
                p = doc.add_paragraph(secao.strip())
                p.style = 'List Number'
            
            # Parágrafo normal
            else:
                # Processar parágrafos
                paragrafos = secao.split('\n\n')
                for para in paragrafos:
                    if para.strip():
                        # Aplicar formatação inline (negrito, itálico)
                        self._adicionar_paragrafo_formatado(doc, para.strip())
    
    def _adicionar_paragrafo_formatado(self, doc, texto):
        """Adiciona parágrafo com formatação inline"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Processar negrito (**texto**)
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                run = p.add_run(parte[2:-2])
                run.bold = True
            else:
                p.add_run(parte)
        
        # Configurar fonte
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    def _adicionar_rodape(self, doc, metadados):
        """Adiciona rodapé com informações da geração"""
        doc.add_page_break()
        
        p = doc.add_paragraph()
        p.add_run("Nestes termos,\nPede deferimento.")
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Assinatura (placeholder)
        assinatura = doc.add_paragraph()
        assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assinatura.add_run("_"*50)
        
        nome_advogado = doc.add_paragraph()
        nome_advogado.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nome_advogado.add_run("[Nome do Advogado]")
        
        oab = doc.add_paragraph()
        oab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        oab.add_run("OAB/RJ nº [número]")
        
        # Informações técnicas (ocultas em fonte pequena)
        doc.add_page_break()
        
        info_tecnica = doc.add_paragraph()
        info_tecnica.add_run("INFORMAÇÕES TÉCNICAS DA GERAÇÃO\n").bold = True
        
        for chave, valor in metadados.items():
            info_tecnica.add_run(f"\n{chave}: {valor}")
        
        for run in info_tecnica.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
